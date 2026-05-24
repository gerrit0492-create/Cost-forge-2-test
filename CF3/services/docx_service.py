from __future__ import annotations

from docx import Document



def generate_docx_quote(path: str, summary: dict):
    doc = Document()

    doc.add_heading('CF3 Enterprise Quote', level=1)

    doc.add_paragraph(f"BOM Lines: {summary.get('bom_lines', 0)}")
    doc.add_paragraph(f"Material Cost: {summary.get('material_cost', 0)}")
    doc.add_paragraph(f"Routing Cost: {summary.get('routing_cost', 0)}")
    doc.add_paragraph(f"Total Cost: {summary.get('total_cost', 0)}")

    doc.save(path)
