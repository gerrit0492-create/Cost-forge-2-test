from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


_STR_COLS = ["line_id", "material_id", "qty"]
_NUM_COLS = ["material_cost", "process_cost", "overhead", "total_cost"]


def _prepare_df(df):
    df = df.copy()
    for col in _STR_COLS:
        if col not in df.columns:
            df[col] = ""
    for col in _NUM_COLS:
        if col not in df.columns:
            df[col] = 0.0
    df[_NUM_COLS] = df[_NUM_COLS].fillna(0.0)
    return df


def make_offer_docx(df, title: str = "Offerte") -> bytes:
    df = _prepare_df(df)
    doc = Document()
    doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.LEFT
    total = float(df["total_cost"].sum())
    p = doc.add_paragraph("Totaalprijs: ")
    p.add_run(f"EUR {total:,.2f}").bold = True

    table = doc.add_table(rows=1, cols=7)
    hdr = ["Line", "Material", "Qty", "Mat. cost", "Proc. cost", "Overhead", "Total"]
    for i, t in enumerate(hdr):
        table.rows[0].cells[i].text = t

    for _, r in df.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["line_id"])
        row[1].text = str(r["material_id"])
        row[2].text = str(r["qty"])
        row[3].text = f"{r['material_cost']:,.2f}"
        row[4].text = f"{r['process_cost']:,.2f}"
        row[5].text = f"{r['overhead']:,.2f}"
        row[6].text = f"{r['total_cost']:,.2f}"

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
